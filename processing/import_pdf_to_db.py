# 📁 processing/import_pdf_to_db.py

import fitz  # PyMuPDF
from docx import Document  # type: ignore  # python-docx (Pylance stubs incomplete)
from pathlib import Path
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import selectinload
from database.db_loader import DBConnector, Source
from datetime import datetime
import logging
from typing import List, Optional, Dict, Tuple, cast
import re
import hashlib
import shutil
import subprocess
import tempfile

class DocumentsToDatabase:
    """Import academic documents (PDF, DOC, DOCX) to database with text extraction and validation"""

    def __init__(self):
        self.db = DBConnector()
        self.session = self.db.get_session()
        self.setup_logging()
        self.last_skip_reason = None

    def setup_logging(self):
        """Setup logging for import tracking"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('document_import_log.txt'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)

    def preprocess_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove common OCR artifacts and headers/footers
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Skip page numbers, dates in headers/footers
            if re.match(r'^\s*-?\s*\d+\s*-?\s*$', line.strip()):
                continue
            # Skip mostly special characters (OCR noise)
            if re.match(r'^[^\w\s]{5,}$', line.strip()):
                continue
            # Skip very short fragments that look like page breaks
            if len(line.strip()) > 0:
                cleaned_lines.append(line)
        
        text = '\n'.join(cleaned_lines)
        
        # Normalize whitespace: remove hyphenated line breaks
        text = re.sub(r'-\n\s*', '', text)
        # Collapse multiple whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text

    def calculate_word_count(self, text: str) -> int:
        """Calculate word count from text"""
        return len(text.split())

    def calculate_reading_time(self, word_count: int) -> int:
        """Calculate estimated reading time in minutes (avg 200 words/minute)"""
        return max(1, word_count // 200)

    def calculate_content_hash(self, text: str) -> str:
        """Generate SHA256 hash of text content for duplicate detection"""
        return hashlib.sha256(text.encode('utf-8')).hexdigest()

    def match_movement(self, text: str) -> Optional[int]:
        """Try to match document content to a known movement by keywords"""
        try:
            # Import movement database
            from database.db_loader import DBConnector
            db = DBConnector()
            session = db.get_session()
            from database.models.movement import Movement
            
            movements = session.query(Movement).options(selectinload(Movement.aliases)).all()
            
            text_lower = text.lower()
            best_match = None
            match_count = 0
            
            # Score each movement based on keyword matches
            for movement in movements:
                movement_name = getattr(movement, "canonical_name", None)
                if not movement_name:
                    continue
                    
                # Check if movement name appears in text
                movement_name = movement_name.lower()
                score = 0
                
                if movement_name in text_lower:
                    score += 10
                
                # Check for movement aliases
                if movement.aliases:
                    for alias in movement.aliases:
                        alias_value = getattr(alias, "alias", None)
                        if alias_value and alias_value.lower() in text_lower:
                            score += 5
                
                if score > match_count:
                    match_count = score
                    best_match = int(movement.id)  # type: ignore  # SQLAlchemy Column type
            
            session.close()
            return best_match
        except Exception as e:
            self.logger.warning(f"Could not match movement by keywords: {e}")
            try:
                session.close()
            except Exception:
                pass
            return None

    def extract_pdf_metadata(self, pdf_path: str) -> Dict:
        """Extract metadata from PDF document properties"""
        metadata = {}
        try:
            doc = fitz.open(pdf_path)
            props = doc.metadata
            
            if props:
                metadata['title'] = props.get('title', '')
                metadata['author'] = props.get('author', '')
                metadata['subject'] = props.get('subject', '')
                metadata['creator'] = props.get('creator', '')
                metadata['producer'] = props.get('producer', '')
                metadata['creation_date'] = props.get('creationDate')
                metadata['mod_date'] = props.get('modDate')
            
            doc.close()
        except Exception as e:
            self.logger.debug(f"Could not extract PDF metadata from {pdf_path}: {e}")
        
        return metadata

    def convert_doc_to_docx(self, doc_path: str) -> Optional[Path]:
        """Convert legacy .doc files to .docx using LibreOffice if available"""
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            self.logger.warning("LibreOffice not found; cannot convert .doc to .docx")
            return None

        temp_dir = Path(tempfile.mkdtemp(prefix="doc_convert_"))
        try:
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(temp_dir),
                    doc_path,
                ],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            out_path = temp_dir / f"{Path(doc_path).stem}.docx"
            if out_path.exists():
                return out_path
            self.logger.warning(f"Converted DOCX not found for {doc_path}")
        except Exception as e:
            self.logger.error(f"Error converting DOC to DOCX {doc_path}: {e}")

        return None

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text content from PDF file with improved preprocessing"""
        try:
            doc = fitz.open(pdf_path)
            text = ""

            # Extract text from all pages
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()  # type: ignore  # PyMuPDF Page.get_text() stubs incomplete
                text += page_text + "\n"

            doc.close()
            
            # Preprocess extracted text
            text = self.preprocess_text(text)
            return text

        except Exception as e:
            self.logger.error(f"Error extracting text from PDF {pdf_path}: {e}")
            return ""

    def extract_text_from_docx(self, doc_path: str) -> str:
        """Extract text content from DOCX/DOC file with improved preprocessing"""
        converted_path: Optional[Path] = None
        temp_dir: Optional[Path] = None
        try:
            path = Path(doc_path)
            if path.suffix.lower() == ".doc":
                converted_path = self.convert_doc_to_docx(doc_path)
                if not converted_path:
                    return ""
                temp_dir = converted_path.parent
                doc_path = str(converted_path)

            doc = Document(doc_path)
            text = ""

            # Extract text from all paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"

            # Also extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"

            # Preprocess extracted text
            text = self.preprocess_text(text)
            return text

        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX/DOC {doc_path}: {e}")
            return ""
        finally:
            if converted_path:
                try:
                    converted_path.unlink()
                except Exception:
                    pass
                if temp_dir:
                    try:
                        temp_dir.rmdir()
                    except Exception:
                        pass

    def extract_text_from_document(self, file_path: str) -> str:
        """Route to appropriate extraction method based on file extension"""
        file_ext = Path(file_path).suffix.lower()

        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        else:
            self.logger.warning(f"Unsupported file format: {file_ext}")
            return ""

    def extract_metadata_from_filename(self, filename: str) -> dict:
        """Extract metadata from document filename"""
        metadata = {
            'title': '',
            'author': '',
            'year': None,
            'type': 'academic_paper'
        }

        # Remove extension and _data suffix
        name = Path(filename).stem.replace('_data', '')

        # Try to extract year
        year_match = re.search(r'_(\d{4})_', filename)
        if year_match:
            metadata['year'] = int(year_match.group(1))

        # Try to extract author
        if '_' in name:
            parts = name.split('_')
            if len(parts) >= 2:
                metadata['author'] = parts[0].replace('_', ' ').title()

        # Create title from filename
        title = name.replace('_', ' ').replace('data', '').strip()
        if metadata['year']:
            title = title.replace(f" {metadata['year']}", '').replace(f"_{metadata['year']}", '')

        # Clean up title
        title = re.sub(r'\s+', ' ', title).strip()
        if not title:
            title = Path(filename).stem

        metadata['title'] = title

        return metadata

    def validate_document_content(self, text: str, file_path: str) -> Tuple[bool, List[str]]:
        """Validate extracted document content comprehensively. Returns (is_valid, error_list)"""
        errors = []
        warnings = []

        if not text or len(text.strip()) < 100:
            errors.append("Extracted text too short or empty")
            return False, errors

        word_count = self.calculate_word_count(text)
        
        # Check minimum word count (at least 50 words)
        if word_count < 50:
            errors.append(f"Text too short ({word_count} words, minimum 50)")
            return False, errors

        # Check for Czech/religious content keywords
        czech_keywords = ['sekta', 'hnutí', 'církev', 'náboženství', 'duchovní', 'religiozní', 'krčem', 'spirituální']
        has_relevant_content = any(keyword in text.lower() for keyword in czech_keywords)

        if not has_relevant_content:
            errors.append("Content doesn't appear to be relevant (missing Czech/religious keywords)")
            return False, errors

        # Check for excessive non-text content (OCR noise)
        non_unicode_ratio = len([c for c in text if ord(c) > 127]) / len(text) if text else 0
        if non_unicode_ratio > 0.3:
            warnings.append(f"High ratio of non-ASCII characters ({non_unicode_ratio:.1%}), possible OCR issues")
        
        # Check for suspicious patterns
        if text.count('\x00') > 0:
            warnings.append("Text contains null bytes (possible corruption)")
        
        self.logger.info(f"Validation: word_count={word_count}, relevant_keywords=found, status=PASS")
        return True, warnings

    def create_source_from_document(self, file_path: str) -> bool:
        """Process single document file and create Source record with enhanced metadata and matching"""
        try:
            doc_file = Path(file_path)
            filename = doc_file.name

            # Extract text
            self.logger.info(f"Processing: {filename}")
            text = self.extract_text_from_document(file_path)

            if not text:
                self.logger.warning(f"No text extracted from {filename}")
                self.logger.info(f"⏭️  Skipped: {filename} (no text extracted)")
                self.last_skip_reason = "no_text"
                return False

            # Calculate content hash early for duplicate detection
            content_hash = self.calculate_content_hash(text)
            
            # Check for duplicates by content hash
            existing_by_hash = self.session.query(Source).filter(Source.content_hash == content_hash).first()
            if existing_by_hash:
                self.logger.info(f"Duplicate content detected (hash match): {filename}")
                self.logger.info(f"⏭️  Skipped: {filename} (duplicate content hash)")
                self.last_skip_reason = "duplicate_content"
                return False

            # Validate content
            is_valid, validation_messages = self.validate_document_content(text, file_path)
            if not is_valid:
                self.logger.warning(f"Validation failed for {filename}: {'; '.join(validation_messages)}")
                self.logger.info(f"⏭️  Skipped: {filename} (validation failed)")
                self.last_skip_reason = "validation_failed"
                return False
            
            for msg in validation_messages:
                self.logger.info(f"  ℹ️  {msg}")

            # Extract metadata
            metadata = self.extract_metadata_from_filename(filename)
            
            # Try to extract metadata from PDF itself
            if doc_file.suffix.lower() == '.pdf':
                pdf_metadata = self.extract_pdf_metadata(file_path)
                if pdf_metadata.get('title'):
                    metadata['title'] = pdf_metadata['title']
                if pdf_metadata.get('author'):
                    metadata['author'] = pdf_metadata['author']
            
            # Calculate analytics
            word_count = self.calculate_word_count(text)
            reading_time = self.calculate_reading_time(word_count)

            # Determine document type from extension
            file_ext = doc_file.suffix.lower()
            source_type = "academic_doc" if file_ext in ['.doc', '.docx'] else "academic_pdf"

            # Try to match to a known movement
            matched_movement_id = self.match_movement(text)
            if matched_movement_id is not None:
                self.logger.info(f"  ✓ Matched to movement ID: {matched_movement_id}")
            else:
                # Get default "Unidentified" movement ID from database
                from database.db_loader import Movement
                default = self.session.query(Movement).filter_by(
                    canonical_name="Neidentifikované hnutí"
                ).first()
                matched_movement_id = cast(int, default.id) if default is not None else None
                if matched_movement_id is not None:
                    self.logger.info(f"  ℹ️  No specific movement match, using default (ID: {matched_movement_id})")
                else:
                    self.logger.warning(f"  ⚠️  No movement match and no default movement available - skipping")
                    return False

            # Create source record
            source = Source(
                movement_id=matched_movement_id,
                source_name=metadata['title'],
                source_type=source_type,
                author=metadata.get('author'),
                url=f"file://{file_path}",  # Local file URL
                content_full=text,
                content_excerpt=text[:500] + "..." if len(text) > 500 else text,
                publication_date=datetime.now(),
                language="cs",  # Assume Czech
                word_count=word_count,
                reading_time_minutes=reading_time,
                content_hash=content_hash,
                scraped_by="pdf_import"
            )

            # Check for duplicates by URL
            existing = self.session.query(Source).filter(Source.url == source.url).first()
            if existing:
                self.logger.info(f"Document already exists in database: {filename}")
                self.logger.info(f"⏭️  Skipped: {filename} (duplicate URL)")
                self.last_skip_reason = "duplicate_url"
                return False

            self.session.add(source)
            self.session.commit()
            self.last_skip_reason = None

            self.logger.info(f"✅ SUCCESS: {filename} ({word_count} words, {reading_time} min read)")
            return True

        except IntegrityError as e:
            self.session.rollback()
            self.logger.error(f"Database integrity error for {file_path}: {e}")
            return False
        except Exception as e:
            self.session.rollback()
            self.logger.error(f"Error processing document {file_path}: {e}", exc_info=True)
            return False

    def load_documents_to_sources(self, documents_directory: str) -> dict:
        """Load all documents (PDF, DOC, DOCX) from directory to database with detailed reporting"""
        stats = {
            'processed': 0,
            'successful': 0,
            'failed': 0,
            'skipped': 0,
            'total_words': 0,
            'total_reading_minutes': 0,
            'skipped_no_text': 0,
            'skipped_duplicate_content': 0,
            'skipped_validation_failed': 0,
            'skipped_duplicate_url': 0,
            'skipped_other': 0
        }

        docs_dir = Path(documents_directory)
        if not docs_dir.exists():
            self.logger.error(f"Documents directory does not exist: {documents_directory}")
            return stats

        # Find all supported document types
        supported_files = []
        supported_files.extend(sorted(docs_dir.glob("*.pdf")))
        supported_files.extend(sorted(docs_dir.glob("*.docx")))
        supported_files.extend(sorted(docs_dir.glob("*.doc")))

        if not supported_files:
            self.logger.warning(f"No supported document files found in {documents_directory}")
            return stats

        self.logger.info(f"\n{'='*60}")
        self.logger.info(f"Starting document import: {len(supported_files)} files found")
        self.logger.info(f"{'='*60}\n")

        for doc_file in supported_files:
            stats['processed'] += 1
            doc_path = str(doc_file)

            try:
                if self.create_source_from_document(doc_path):
                    stats['successful'] += 1
                    # Get word count from database
                    source = self.session.query(Source).filter(Source.url == f"file://{doc_path}").first()
                    if source:
                        stats['total_words'] += int(source.word_count or 0)  # type: ignore
                        stats['total_reading_minutes'] += int(source.reading_time_minutes or 0)  # type: ignore
                else:
                    stats['skipped'] += 1
                    reason = self.last_skip_reason or "other"
                    if reason == "no_text":
                        stats['skipped_no_text'] += 1
                    elif reason == "duplicate_content":
                        stats['skipped_duplicate_content'] += 1
                    elif reason == "validation_failed":
                        stats['skipped_validation_failed'] += 1
                    elif reason == "duplicate_url":
                        stats['skipped_duplicate_url'] += 1
                    else:
                        stats['skipped_other'] += 1
            except Exception as e:
                stats['failed'] += 1
                self.logger.error(f"FAILED: {doc_file.name}: {e}")
                continue

        self.session.close()
        
        # Print summary
        self.logger.info(f"\n{'='*60}")
        self.logger.info(f"Import Summary:")
        self.logger.info(f"  Total processed:     {stats['processed']}")
        self.logger.info(f"  ✅ Successful:       {stats['successful']}")
        self.logger.info(f"  ⏭️  Skipped:         {stats['skipped']}")
        self.logger.info(f"     • No text:        {stats['skipped_no_text']}")
        self.logger.info(f"     • Duplicate hash: {stats['skipped_duplicate_content']}")
        self.logger.info(f"     • Validation:     {stats['skipped_validation_failed']}")
        self.logger.info(f"     • Duplicate URL:  {stats['skipped_duplicate_url']}")
        self.logger.info(f"     • Other:          {stats['skipped_other']}")
        self.logger.info(f"  ❌ Failed:           {stats['failed']}")
        self.logger.info(f"  Total words:        {stats['total_words']:,}")
        self.logger.info(f"  Total reading time: {stats['total_reading_minutes']} minutes")
        self.logger.info(f"{'='*60}\n")
        
        return stats


# Backward compatibility alias
class PDFtoDatabaseLoader(DocumentsToDatabase):
    """Backward compatibility wrapper for old class name"""
    
    def load_pdfs_to_sources(self, pdf_directory: str) -> dict:
        """Backward compatibility method"""
        return self.load_documents_to_sources(pdf_directory)
